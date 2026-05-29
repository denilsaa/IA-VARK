import json
import os
import re
import unicodedata

from django.utils.text import Truncator

from .models import MaterialEstudio


MAX_CARACTERES_EXTRAIDOS = 300000
MAX_CARACTERES_PARA_IA = 45000
MAX_FRAGMENTOS_RELEVANTES = 12


def procesar_material(material):
    """
    Extrae texto del material y luego lo analiza con Gemini.
    """

    material.estado = MaterialEstudio.ESTADO_PROCESANDO
    material.error_procesamiento = ""
    material.save(update_fields=["estado", "error_procesamiento", "actualizado"])

    try:
        texto = obtener_texto_material(material)
        texto = limpiar_texto(texto)

        if not texto:
            material.estado = MaterialEstudio.ESTADO_ERROR
            material.error_procesamiento = "No se pudo extraer texto del material."
            material.save(
                update_fields=[
                    "estado",
                    "error_procesamiento",
                    "actualizado",
                ]
            )
            return material

        material.texto_extraido = texto[:MAX_CARACTERES_EXTRAIDOS]

        analisis = analizar_texto_con_gemini(material, material.texto_extraido)

        if analisis:
            material.resumen_ia = analisis.get("resumen", "")
            material.temas_clave_ia = "\n".join(analisis.get("temas_clave", []))
            material.preguntas_sugeridas_ia = "\n".join(
                analisis.get("preguntas_sugeridas", [])
            )
            material.recomendacion_ia = analisis.get("recomendacion", "")

        material.estado = MaterialEstudio.ESTADO_PROCESADO
        material.save(
            update_fields=[
                "texto_extraido",
                "resumen_ia",
                "temas_clave_ia",
                "preguntas_sugeridas_ia",
                "recomendacion_ia",
                "estado",
                "actualizado",
            ]
        )

        return material

    except Exception as error:
        material.estado = MaterialEstudio.ESTADO_ERROR
        material.error_procesamiento = str(error)
        material.save(
            update_fields=[
                "estado",
                "error_procesamiento",
                "actualizado",
            ]
        )
        return material


def obtener_texto_material(material):
    if material.texto_manual and material.texto_manual.strip():
        return material.texto_manual.strip()

    if not material.archivo:
        return ""

    ruta = material.archivo.path
    extension = os.path.splitext(ruta)[1].lower()

    if extension == ".pdf":
        return extraer_texto_pdf(ruta)

    if extension == ".docx":
        return extraer_texto_docx(ruta)

    if extension == ".txt":
        return extraer_texto_txt(ruta)

    if extension in [".jpg", ".jpeg", ".png", ".webp"]:
        return extraer_texto_imagen(ruta)

    if extension == ".doc":
        return (
            "El archivo .doc fue guardado, pero este formato antiguo no se procesa "
            "en esta versión. Convierte el documento a .docx o PDF."
        )

    return ""


def extraer_texto_pdf(ruta):
    import fitz

    partes = []

    with fitz.open(ruta) as documento:
        total_paginas = len(documento)

        for numero_pagina in range(total_paginas):
            pagina = documento.load_page(numero_pagina)
            texto = pagina.get_text("text")

            if texto.strip():
                partes.append(f"\n--- Página {numero_pagina + 1} ---\n{texto}")

            if len("\n".join(partes)) >= MAX_CARACTERES_EXTRAIDOS:
                break

    return "\n".join(partes)


def extraer_texto_docx(ruta):
    from docx import Document

    documento = Document(ruta)
    partes = []

    for parrafo in documento.paragraphs:
        texto = parrafo.text.strip()

        if texto:
            partes.append(texto)

    for tabla in documento.tables:
        for fila in tabla.rows:
            celdas = [
                celda.text.strip()
                for celda in fila.cells
                if celda.text.strip()
            ]

            if celdas:
                partes.append(" | ".join(celdas))

    return "\n".join(partes)


def extraer_texto_txt(ruta):
    codificaciones = ["utf-8", "latin-1", "cp1252"]

    for codificacion in codificaciones:
        try:
            with open(ruta, "r", encoding=codificacion) as archivo:
                return archivo.read()
        except UnicodeDecodeError:
            continue

    return ""


def extraer_texto_imagen(ruta):
    from PIL import Image
    import pytesseract

    imagen = Image.open(ruta)

    try:
        return pytesseract.image_to_string(imagen, lang="spa+eng")
    except Exception:
        return pytesseract.image_to_string(imagen)


def limpiar_texto(texto):
    texto = str(texto or "")
    lineas = [linea.strip() for linea in texto.splitlines()]
    lineas = [linea for linea in lineas if linea]
    return "\n".join(lineas)


def analizar_texto_con_gemini(material, texto):
    api_key = os.getenv("GEMINI_API_KEY", "").strip()

    if not api_key:
        return {}

    if os.getenv("LLM_PROVIDER", "gemini").strip().lower() != "gemini":
        return {}

    try:
        from google import genai
        from google.genai import types

        client = genai.Client(api_key=api_key)
        model = os.getenv("GEMINI_MODEL", "gemini-2.5-flash-lite").strip()

        texto_relevante = seleccionar_texto_relevante(
            texto=texto,
            temario=material.temario_examen,
        )

        texto_para_ia = Truncator(texto_relevante).chars(MAX_CARACTERES_PARA_IA)

        prompt = f"""
Analiza el siguiente material de estudio de Anatomía I.

Datos del material:
- Título: {material.titulo}
- Tema general: {material.tema or "No especificado"}
- Descripción: {material.descripcion or "Sin descripción"}

Temas que entrarán al examen:
\"\"\"
{material.temario_examen or "No se especificaron temas concretos."}
\"\"\"

Texto relevante extraído del material:
\"\"\"
{texto_para_ia}
\"\"\"

Devuelve únicamente JSON válido con esta estructura exacta:

{{
  "resumen": "Resumen breve del material enfocado en los temas del examen, máximo 180 palabras.",
  "temas_clave": [
    "Tema clave 1",
    "Tema clave 2",
    "Tema clave 3"
  ],
  "preguntas_sugeridas": [
    "Pregunta de práctica 1",
    "Pregunta de práctica 2",
    "Pregunta de práctica 3"
  ],
  "recomendacion": "Recomendación breve sobre cómo estudiar este material para el examen."
}}

Reglas:
- Escribe en español.
- Enfócate en los temas que entrarán al examen.
- Si el texto relevante no cubre todos los temas del examen, dilo en la recomendación.
- No inventes contenido que no aparezca o no esté relacionado con el texto.
- No uses markdown.
- No agregues texto fuera del JSON.
"""

        response = client.models.generate_content(
            model=model,
            contents=prompt,
            config=types.GenerateContentConfig(
                response_mime_type="application/json",
                temperature=0.4,
            ),
        )

        contenido = limpiar_json(response.text)
        data = json.loads(contenido)

        return {
            "resumen": str(data.get("resumen", "")).strip(),
            "temas_clave": normalizar_lista(data.get("temas_clave", [])),
            "preguntas_sugeridas": normalizar_lista(
                data.get("preguntas_sugeridas", [])
            ),
            "recomendacion": str(data.get("recomendacion", "")).strip(),
        }

    except Exception as error:
        print("Error analizando material con Gemini:", error)
        return {}


def seleccionar_texto_relevante(texto, temario):
    """
    Para libros grandes, selecciona páginas/fragmentos relacionados con los temas del examen.
    Si no encuentra coincidencias, usa el inicio del texto.
    """

    texto = str(texto or "")
    temario = str(temario or "").strip()

    if not temario:
        return texto[:MAX_CARACTERES_PARA_IA]

    fragmentos = dividir_por_paginas(texto)
    consultas = extraer_consultas_temario(temario)

    if not consultas:
        return texto[:MAX_CARACTERES_PARA_IA]

    fragmentos_puntuados = []

    for etiqueta, contenido in fragmentos:
        contenido_normalizado = normalizar_texto(contenido)
        puntaje = calcular_puntaje_fragmento(contenido_normalizado, consultas)

        if puntaje > 0:
            fragmentos_puntuados.append(
                {
                    "puntaje": puntaje,
                    "etiqueta": etiqueta,
                    "contenido": contenido,
                }
            )

    if not fragmentos_puntuados:
        return texto[:MAX_CARACTERES_PARA_IA]

    fragmentos_puntuados.sort(key=lambda item: item["puntaje"], reverse=True)
    seleccionados = fragmentos_puntuados[:MAX_FRAGMENTOS_RELEVANTES]

    partes = []

    for item in seleccionados:
        partes.append(f"\n--- {item['etiqueta']} ---\n{item['contenido']}")

    return "\n".join(partes)


def dividir_por_paginas(texto):
    patron = re.compile(r"--- Página\s+(\d+)\s+---", re.IGNORECASE)
    coincidencias = list(patron.finditer(texto))

    if not coincidencias:
        return [("Fragmento 1", texto)]

    fragmentos = []

    for index, coincidencia in enumerate(coincidencias):
        numero_pagina = coincidencia.group(1)
        inicio = coincidencia.end()

        if index + 1 < len(coincidencias):
            fin = coincidencias[index + 1].start()
        else:
            fin = len(texto)

        contenido = texto[inicio:fin].strip()

        if contenido:
            fragmentos.append((f"Página {numero_pagina}", contenido))

    return fragmentos


def extraer_consultas_temario(temario):
    lineas = re.split(r"[\n,;]+", temario)
    consultas = []

    for linea in lineas:
        linea = linea.strip()

        if not linea:
            continue

        frase = normalizar_texto(linea)
        palabras = [
            palabra
            for palabra in re.findall(r"[a-záéíóúñü]+", frase)
            if len(palabra) >= 4
        ]

        consultas.append(
            {
                "frase": frase,
                "palabras": palabras,
            }
        )

    return consultas


def calcular_puntaje_fragmento(contenido_normalizado, consultas):
    puntaje = 0

    for consulta in consultas:
        frase = consulta["frase"]
        palabras = consulta["palabras"]

        if frase and frase in contenido_normalizado:
            puntaje += 10

        for palabra in palabras:
            if palabra in contenido_normalizado:
                puntaje += 1

    return puntaje


def normalizar_texto(texto):
    texto = str(texto or "").lower()
    texto = unicodedata.normalize("NFD", texto)
    texto = "".join(
        caracter
        for caracter in texto
        if unicodedata.category(caracter) != "Mn"
    )
    texto = texto.replace("ñ", "n")
    texto = re.sub(r"\s+", " ", texto)
    return texto.strip()


def normalizar_lista(valor):
    if isinstance(valor, list):
        return [str(item).strip() for item in valor if str(item).strip()]

    if isinstance(valor, str):
        return [linea.strip() for linea in valor.splitlines() if linea.strip()]

    return []


def limpiar_json(texto):
    texto = str(texto or "").strip()

    if texto.startswith("```json"):
        texto = texto.replace("```json", "", 1).strip()

    if texto.startswith("```"):
        texto = texto.replace("```", "", 1).strip()

    if texto.endswith("```"):
        texto = texto[:-3].strip()

    return texto